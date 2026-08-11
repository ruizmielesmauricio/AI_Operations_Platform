import io
from unittest.mock import patch

import pytest
from docx import Document
from reportlab.graphics.shapes import Drawing, String
from reportlab.platypus import Paragraph, Table

from app.exports.docx import render_report_docx
from app.exports.pdf import render_report_pdf


def complete_report_payload(report_type: str = "monthly") -> dict:
    return {
        "business_name": "Complete Cycle Shop",
        "business_type": "bicycle_shop",
        "report_type": report_type,
        "period_start": "2026-07-01T00:00:00Z",
        "period_end": "2026-07-31T23:59:59Z",
        "generated_at": "2026-08-01T08:00:00Z",
        "executive_summary": {
            "narrative": ["Revenue was steady and inventory needs attention."],
            "transactions": 42,
            "average_sale": "52.50",
            "low_stock_count": 2,
            "dead_stock_count": 1,
            "top_recommendations": [
                {
                    "finding_type": "dead_stock",
                    "severity": "warning",
                    "title": "Review dead stock",
                    "description": "Consider a bundle or supplier return.",
                    "evidence": {"product_id": "p-dead", "category_name": "Accessories"},
                    "impact_score": "8.0",
                }
            ],
        },
        "financial_performance": {
            "period": {"start": "2026-07-01", "end": "2026-07-31"},
            "revenue": {"current": "2205.00", "previous": "2100.00", "change_pct": "5.00"},
            "gross_margin": {
                "total_revenue": "2205.00",
                "revenue_with_known_cost": "2100.00",
                "cogs": "1200.00",
                "gross_profit": "900.00",
                "gross_margin_pct": "42.86",
                "cost_data_coverage_pct": "95.24",
                "net_gross_profit": "850.00",
                "net_gross_margin_pct": "40.48",
                "tax_data_coverage_pct": "90.00",
            },
            "top_margin_products": [
                {
                    "product_id": "p-helmet",
                    "name": "Road Helmet",
                    "revenue": "600.00",
                    "gross_profit": "280.00",
                    "gross_margin_pct": "46.67",
                    "category_name": "Helmets",
                }
            ],
            "bottom_margin_products": [
                {
                    "product_id": "p-tube",
                    "name": "Inner Tube",
                    "revenue": "200.00",
                    "gross_profit": "30.00",
                    "gross_margin_pct": "15.00",
                    "category_name": "Components",
                }
            ],
            "products_excluded_from_ranking": 1,
            "returns": {
                "gross_revenue": "2250.00",
                "returns_amount": "45.00",
                "return_count": 1,
                "net_revenue": "2205.00",
                "return_rate_pct": "2.00",
            },
        },
        "retail_operations": {
            "period": {"start": "2026-07-01", "end": "2026-07-31"},
            "top_sellers_by_units": [
                {
                    "product_id": "p-tube",
                    "name": "Inner Tube",
                    "units_sold": 20,
                    "revenue": "200.00",
                    "category_name": "Components",
                }
            ],
            "top_sellers_by_revenue": [
                {
                    "product_id": "p-helmet",
                    "name": "Road Helmet",
                    "units_sold": 8,
                    "revenue": "600.00",
                    "category_name": "Helmets",
                }
            ],
            "stock_cover": [
                {
                    "product_id": "p-helmet",
                    "name": "Road Helmet",
                    "stock_on_hand": 12,
                    "units_sold_in_period": 8,
                    "cover_days": "46.50",
                    "revenue_in_period": "600.00",
                    "category_name": "Helmets",
                }
            ],
            "dead_stock": [
                {
                    "product_id": "p-dead",
                    "name": "Old Saddle",
                    "stock_on_hand": 4,
                    "value_at_cost": "120.00",
                    "category_name": "Accessories",
                }
            ],
            "inventory_value": {"value_at_cost": "4200.00", "products_missing_cost": 2},
            "sell_through_rate": "0.315",
        },
        "inventory_health": {
            "fast_movers": [
                {
                    "product_id": "p-fast",
                    "name": "Brake Pads",
                    "stock_on_hand": 3,
                    "units_sold_in_period": 12,
                    "cover_days": "7.75",
                    "revenue_in_period": "240.00",
                    "category_name": "Components",
                }
            ],
            "slow_movers": [
                {
                    "product_id": "p-slow",
                    "name": "Touring Rack",
                    "stock_on_hand": 10,
                    "units_sold_in_period": 1,
                    "cover_days": "310.00",
                    "revenue_in_period": "85.00",
                    "category_name": "Accessories",
                }
            ],
            "turnover_ratio": "1.80",
        },
        "forecast": {
            "horizon_days": 2,
            "revenue": {
                "horizon_days": 2,
                "result": {
                    "insufficient_data": False,
                    "method": "seasonal_day_of_week",
                    "history_days_used": 90,
                    "daily": [
                        {"forecast_date": "2026-08-01", "point": "100.00", "low": "80.00", "high": "120.00"},
                        {"forecast_date": "2026-08-02", "point": "130.00", "low": "100.00", "high": "155.00"},
                    ],
                    "total_point": "230.00",
                    "total_low": "180.00",
                    "total_high": "275.00",
                },
            },
            "products": [
                {
                    "product_id": "p-fast",
                    "name": "Brake Pads",
                    "sku": "BP-01",
                    "result": {
                        "insufficient_data": False,
                        "method": "moving_average",
                        "history_days_used": 90,
                        "daily": [],
                        "total_point": "6",
                        "total_low": "4",
                        "total_high": "9",
                    },
                    "current_stock": 3,
                    "suggested_reorder_quantity": 6,
                    "days_of_cover_at_forecast_rate": "7.00",
                    "category_name": "Components",
                }
            ],
            "products_excluded_insufficient_data": 2,
        },
        "findings": {
            "period": {"start": "2026-07-01", "end": "2026-07-31"},
            "findings": [
                {
                    "type": "dead_stock",
                    "severity": "warning",
                    "message": "Old Saddle had no sales this month.",
                    "evidence": {"product_id": "p-dead", "category_name": "Accessories"},
                    "rule_id": "dead-stock",
                    "rule_version": 1,
                },
                {
                    "type": "low_gross_margin",
                    "severity": "warning",
                    "message": "Margin is below the target range.",
                    "evidence": {},
                    "rule_id": "margin",
                    "rule_version": 1,
                },
            ],
            "recommendations": [
                {
                    "finding_type": "low_gross_margin",
                    "severity": "warning",
                    "title": "Review pricing",
                    "description": "Check high-volume products with weak margin.",
                    "evidence": {},
                    "impact_score": "7.0",
                },
                {
                    "finding_type": "dead_stock",
                    "severity": "warning",
                    "title": "Review dead stock",
                    "description": "Consider a bundle or supplier return.",
                    "evidence": {"product_id": "p-dead", "category_name": "Accessories"},
                    "impact_score": "8.0",
                },
            ],
        },
        "workshop_performance": {
            "period": {"start": "2026-07-01", "end": "2026-07-31"},
            "revenue": {"current": "800.00", "previous": "700.00", "change_pct": "14.29"},
            "margin": {
                "repair_count": 10,
                "revenue": "800.00",
                "revenue_coverage_pct": "100.00",
                "labour_cost": "300.00",
                "gross_profit": "500.00",
                "gross_margin_pct": "62.50",
                "labour_cost_coverage_pct": "100.00",
                "average_ticket": "80.00",
                "net_gross_profit": "450.00",
                "net_gross_margin_pct": "60.00",
                "tax_data_coverage_pct": "90.00",
            },
        },
        "category_breakdown": {
            "period": {"start": "2026-07-01", "end": "2026-07-31"},
            "rows": [
                {
                    "category_id": "cat-helmets",
                    "category_name": "Helmets",
                    "revenue": "600.00",
                    "expenses": "210.00",
                    "expenses_data_coverage_pct": "75.00",
                    "stock_value": "1800.00",
                    "products_excluded_from_stock_value": 1,
                },
                {
                    "category_id": None,
                    "category_name": "Uncategorized",
                    "revenue": "100.00",
                    "expenses": "0.00",
                    "expenses_data_coverage_pct": None,
                    "stock_value": "250.00",
                    "products_excluded_from_stock_value": 0,
                },
            ],
        },
    }


@pytest.mark.parametrize("report_type,label", [("weekly", "Weekly Report"), ("monthly", "Monthly Report")])
def test_pdf_contains_complete_weekly_and_monthly_report(report_type, label):
    with patch("app.exports.pdf.SimpleDocTemplate.build") as build:
        render_report_pdf(complete_report_payload(report_type))
    story = build.call_args.args[0]

    text_parts = []
    for flowable in story:
        if isinstance(flowable, Paragraph):
            text_parts.append(flowable.getPlainText())
        elif isinstance(flowable, Table):
            for row in flowable._cellvalues:
                for cell in row:
                    if isinstance(cell, Paragraph):
                        text_parts.append(cell.getPlainText())
        elif isinstance(flowable, Drawing):
            text_parts.extend(node.text for node in flowable.contents if isinstance(node, String))
    text = "\n".join(text_parts)

    assert label in text
    for expected in (
        "Category Breakdown",
        "Helmets",
        "Gross profit by product",
        "Estimated days of stock cover",
        "Daily revenue forecast",
        "Fast movers",
        "Slow movers",
        "Old Saddle",
        "Purchasing Recommendations",
        "Workshop Performance",
        "Business performance",
        "Stock & products",
    ):
        assert expected in text


def test_pdf_renderer_returns_a_real_pdf():
    exported = render_report_pdf(complete_report_payload())
    assert exported.startswith(b"%PDF")
    assert len(exported) > 10_000


@pytest.mark.parametrize("report_type,label", [("weekly", "Weekly Report"), ("monthly", "Monthly Report")])
def test_docx_contains_complete_weekly_and_monthly_report(report_type, label):
    exported = render_report_docx(complete_report_payload(report_type))
    document = Document(io.BytesIO(exported))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    combined = f"{text}\n{table_text}"

    assert label in combined
    for expected in (
        "Category Breakdown",
        "Helmets",
        "Gross profit by product",
        "Daily revenue forecast",
        "Fast movers",
        "Slow movers",
        "Old Saddle",
        "Purchasing Recommendations",
        "Workshop Performance",
        "Business performance",
        "Stock & products",
    ):
        assert expected in combined
