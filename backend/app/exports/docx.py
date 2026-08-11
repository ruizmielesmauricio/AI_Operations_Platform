"""Complete DOCX rendering for stored weekly and monthly reports.

The PDF carries the visual charts. DOCX includes every section and the full
chart-data tables so no report information is lost in an editable format.
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.exports.formatting import format_money, format_pct, format_rate, safe_get, safe_str

_BUSINESS_WIDE_FINDING_TYPES = {
    "revenue_decline",
    "low_gross_margin",
    "incomplete_cost_data",
    "high_return_rate",
}


def _shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_cell_text(cell, value: object, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(safe_str(value))
    run.bold = bold
    run.font.size = Pt(8)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_table(document: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header, bold=True, color="FFFFFF")
        _shade_cell(table.rows[0].cells[index], "17212B")
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            _set_cell_text(cells[index], value)
            if row_index % 2:
                _shade_cell(cells[index], "F4F6F8")
    document.add_paragraph()


def _add_metrics(document: Document, rows: list[tuple[str, object]]) -> None:
    _add_table(document, ["Metric", "Value"], [[label, value] for label, value in rows])


def _product_name(row: dict) -> str:
    name = safe_str(row.get("name"))
    category = row.get("category_name")
    return f"{name} - {category}" if category else name


def _add_recommendations(document: Document, recommendations: list[dict], findings: list[dict]) -> None:
    finding_by_key = {
        (finding.get("type"), (finding.get("evidence") or {}).get("product_id")): finding
        for finding in findings
    }
    for recommendation in recommendations:
        evidence = recommendation.get("evidence") or {}
        finding = finding_by_key.get((recommendation.get("finding_type"), evidence.get("product_id")))
        category = evidence.get("category_name")
        title = safe_str(recommendation.get("title"))
        if category:
            title = f"{title} - {category}"
        paragraph = document.add_paragraph(style="List Bullet")
        heading = paragraph.add_run(f"{safe_str(recommendation.get('severity'), 'info').upper()}: {title}")
        heading.bold = True
        if finding:
            paragraph.add_run(f" - {safe_str(finding.get('message'))}")
        paragraph.add_run(f"\n{safe_str(recommendation.get('description'), '')}")


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Arial"
        styles[style_name].font.color.rgb = RGBColor(23, 33, 43)
    footer = section.footer.paragraphs[0]
    footer.text = "ORLA report - deterministic business analytics"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(95, 107, 118)


def render_report_docx(payload: dict) -> bytes:
    document = Document()
    _configure_document(document)

    label = "Weekly" if payload.get("report_type") == "weekly" else "Monthly"
    title = document.add_heading(f"{safe_str(payload.get('business_name'))} - {label} Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph(
        f"{safe_str(payload.get('period_start'))} to {safe_str(payload.get('period_end'))} | "
        f"Generated {safe_str(payload.get('generated_at'))}"
    )
    note = document.add_paragraph(
        "Every figure in this export comes from the same deterministic calculations used by the ORLA report page."
    )
    note.style = document.styles["Quote"]

    summary = safe_get(payload, "executive_summary", default={})
    financial = safe_get(payload, "financial_performance", default={})
    retail = safe_get(payload, "retail_operations", default={})
    inventory = safe_get(payload, "inventory_health", default={})
    forecast = safe_get(payload, "forecast", default={})
    findings_section = safe_get(payload, "findings", default={})
    findings = findings_section.get("findings") or []
    recommendations = findings_section.get("recommendations") or []

    document.add_heading("Executive Summary", level=2)
    for line in summary.get("narrative") or []:
        document.add_paragraph(safe_str(line), style="List Bullet")
    gross_margin = financial.get("gross_margin") or {}
    preferred_margin_pct = gross_margin.get("net_gross_margin_pct") or gross_margin.get("gross_margin_pct")
    _add_metrics(document, [
        ("Revenue", format_money(safe_get(financial, "revenue", "current"))),
        ("Transactions", summary.get("transactions")),
        ("Average sale", format_money(summary.get("average_sale"))),
        ("Gross margin", format_pct(preferred_margin_pct)),
        ("Inventory value at cost", format_money(safe_get(retail, "inventory_value", "value_at_cost"))),
        ("Low stock products", summary.get("low_stock_count")),
        ("Dead stock products", summary.get("dead_stock_count")),
    ])
    top_recommendations = summary.get("top_recommendations") or []
    if top_recommendations:
        document.add_heading("Top recommendations", level=3)
        _add_recommendations(document, top_recommendations, findings)

    document.add_heading("Revenue Performance", level=2)
    revenue = financial.get("revenue") or {}
    trend_name = "last week" if payload.get("report_type") == "weekly" else "last month"
    _add_metrics(document, [
        ("Revenue", format_money(revenue.get("current"))),
        (f"Revenue - {trend_name}", format_money(revenue.get("previous"))),
        (f"Change vs {trend_name}", format_pct(revenue.get("change_pct"))),
        ("Gross profit", format_money(gross_margin.get("net_gross_profit") or gross_margin.get("gross_profit"))),
        ("Gross margin", format_pct(preferred_margin_pct)),
        ("Cost data coverage", format_pct(gross_margin.get("cost_data_coverage_pct"))),
        ("Tax data coverage", format_pct(gross_margin.get("tax_data_coverage_pct"))),
    ])
    returns = financial.get("returns") or {}
    if returns and str(returns.get("returns_amount", "0")) not in {"0", "0.0", "0.00"}:
        document.add_heading("Returns", level=3)
        _add_metrics(document, [
            ("Gross revenue before returns", format_money(returns.get("gross_revenue"))),
            ("Return count", returns.get("return_count")),
            ("Returns amount", format_money(returns.get("returns_amount"))),
            ("Return rate", format_pct(returns.get("return_rate_pct"))),
            ("Net revenue", format_money(returns.get("net_revenue"))),
        ])

    margin_rows_by_id = {}
    for row in [*(financial.get("bottom_margin_products") or []), *(financial.get("top_margin_products") or [])]:
        margin_rows_by_id[row.get("product_id") or row.get("name")] = row
    margin_rows = list(margin_rows_by_id.values())
    if margin_rows:
        document.add_heading("Gross profit by product", level=3)
        _add_table(
            document,
            ["Product", "Revenue", "Gross profit", "Margin"],
            [[_product_name(row), format_money(row.get("revenue")), format_money(row.get("gross_profit")), format_pct(row.get("gross_margin_pct"))] for row in margin_rows],
        )
    if financial.get("products_excluded_from_ranking"):
        document.add_paragraph(
            f"{financial['products_excluded_from_ranking']} product(s) were excluded because cost data was unavailable."
        )

    document.add_heading("Sales Performance", level=2)
    for heading, rows in (
        ("Top sellers by units", retail.get("top_sellers_by_units") or []),
        ("Top sellers by revenue", retail.get("top_sellers_by_revenue") or []),
    ):
        document.add_heading(heading, level=3)
        if rows:
            _add_table(
                document,
                ["Product", "Units sold", "Revenue"],
                [[_product_name(row), row.get("units_sold"), format_money(row.get("revenue"))] for row in rows],
            )
        else:
            document.add_paragraph("No sales in this period.")

    document.add_heading("Category Breakdown", level=2)
    document.add_paragraph(
        "Expenses are purchase cost, not cost of goods sold. Stock value is shown at sell price. "
        "Products without a category are grouped as Uncategorized."
    )
    category_rows = safe_get(payload, "category_breakdown", "rows", default=[])
    if category_rows:
        _add_table(
            document,
            ["Category", "Revenue", "Expenses", "Stock value", "Data notes"],
            [
                [
                    row.get("category_name"),
                    format_money(row.get("revenue")),
                    format_money(row.get("expenses")),
                    format_money(row.get("stock_value")),
                    "; ".join(
                        note for note in [
                            f"Purchase cost coverage {format_pct(row.get('expenses_data_coverage_pct'))}"
                            if row.get("expenses_data_coverage_pct") is not None
                            and float(row.get("expenses_data_coverage_pct")) < 100
                            else "",
                            f"{row.get('products_excluded_from_stock_value')} product(s) excluded - no sell price"
                            if row.get("products_excluded_from_stock_value")
                            else "",
                        ] if note
                    ) or "-",
                ]
                for row in category_rows
            ],
        )
    else:
        document.add_paragraph("No category data for this period.")

    document.add_page_break()
    document.add_heading("Inventory Health", level=2)
    _add_metrics(document, [
        ("Inventory value at cost", format_money(safe_get(retail, "inventory_value", "value_at_cost"))),
        ("Products missing cost", safe_get(retail, "inventory_value", "products_missing_cost")),
        ("Sell-through rate", format_rate(retail.get("sell_through_rate"))),
        ("Inventory turnover", f"{inventory.get('turnover_ratio')}x" if inventory.get("turnover_ratio") is not None else "-"),
    ])
    stock_cover = [row for row in (retail.get("stock_cover") or []) if row.get("cover_days") is not None]
    document.add_heading("Stock cover", level=3)
    if stock_cover:
        _add_table(
            document,
            ["Product", "In stock", "Sold", "Cover", "Revenue"],
            [[_product_name(row), row.get("stock_on_hand"), row.get("units_sold_in_period"), f"{row.get('cover_days')} days", format_money(row.get("revenue_in_period"))] for row in stock_cover],
        )
    else:
        document.add_paragraph("Not enough recent sales to estimate stock cover.")

    for heading, explanation, rows in (
        ("Fast movers", "Products selling through quickly (14 days of cover or less).", inventory.get("fast_movers") or []),
        ("Slow movers", "Products with 60 or more days of cover.", inventory.get("slow_movers") or []),
    ):
        document.add_heading(heading, level=3)
        document.add_paragraph(explanation)
        if rows:
            _add_table(
                document,
                ["Product", "Stock on hand", "Cover left"],
                [[_product_name(row), row.get("stock_on_hand"), f"{row.get('cover_days')} days"] for row in rows],
            )
        else:
            document.add_paragraph("None this period.")

    document.add_heading("Dead stock", level=3)
    dead_stock = retail.get("dead_stock") or []
    if dead_stock:
        _add_table(
            document,
            ["Product", "Stock on hand", "Value at cost"],
            [[_product_name(row), row.get("stock_on_hand"), format_money(row.get("value_at_cost")) if row.get("value_at_cost") is not None else "Unknown"] for row in dead_stock],
        )
    else:
        document.add_paragraph("None - every product with stock on hand sold at least once this period.")

    document.add_heading("Forecast & Future Outlook", level=2)
    revenue_forecast = safe_get(forecast, "revenue", "result", default={})
    if revenue_forecast.get("insufficient_data"):
        document.add_paragraph("Not enough sales history yet to forecast revenue.")
    else:
        horizon = forecast.get("horizon_days") or safe_get(forecast, "revenue", "horizon_days")
        _add_metrics(document, [
            (f"Expected revenue - next {horizon} days", format_money(revenue_forecast.get("total_point"))),
            ("Expected range", f"{format_money(revenue_forecast.get('total_low'))} to {format_money(revenue_forecast.get('total_high'))}"),
            ("Forecast method", revenue_forecast.get("method")),
            ("History used", f"{revenue_forecast.get('history_days_used')} days"),
        ])
        daily = revenue_forecast.get("daily") or []
        if daily:
            document.add_heading("Daily revenue forecast", level=3)
            _add_table(
                document,
                ["Date", "Expected", "Low", "High"],
                [[row.get("forecast_date"), format_money(row.get("point")), format_money(row.get("low")), format_money(row.get("high"))] for row in daily],
            )

    document.add_heading("Purchasing Recommendations", level=2)
    product_forecasts = forecast.get("products") or []
    if product_forecasts:
        _add_table(
            document,
            ["Product", "Current", "Forecast demand", "Cover", "Reorder"],
            [
                [
                    _product_name(row),
                    row.get("current_stock"),
                    f"{safe_str(safe_get(row, 'result', 'total_point'))} ({safe_str(safe_get(row, 'result', 'total_low'))}-{safe_str(safe_get(row, 'result', 'total_high'))})",
                    f"{row.get('days_of_cover_at_forecast_rate')} days" if row.get("days_of_cover_at_forecast_rate") is not None else "-",
                    row.get("suggested_reorder_quantity") or "-",
                ]
                for row in product_forecasts
            ],
        )
        if forecast.get("products_excluded_insufficient_data"):
            document.add_paragraph(
                f"{forecast['products_excluded_insufficient_data']} product(s) were excluded because there was not enough sales history."
            )
    else:
        document.add_paragraph("No products have enough sales history yet to forecast demand.")

    workshop = payload.get("workshop_performance")
    if workshop:
        workshop_margin = workshop.get("margin") or {}
        workshop_revenue = workshop.get("revenue") or {}
        workshop_margin_pct = workshop_margin.get("net_gross_margin_pct") or workshop_margin.get("gross_margin_pct")
        document.add_heading("Workshop Performance", level=2)
        _add_metrics(document, [
            ("Repairs completed", workshop_margin.get("repair_count")),
            ("Revenue", format_money(workshop_revenue.get("current"))),
            (f"Revenue - {trend_name}", format_money(workshop_revenue.get("previous"))),
            (f"Change vs {trend_name}", format_pct(workshop_revenue.get("change_pct"))),
            ("Average ticket", format_money(workshop_margin.get("average_ticket"))),
            ("Gross margin - labour only", format_pct(workshop_margin_pct)),
            ("Labour cost coverage", format_pct(workshop_margin.get("labour_cost_coverage_pct"))),
            ("Tax data coverage", format_pct(workshop_margin.get("tax_data_coverage_pct"))),
        ])

    document.add_heading("Action Plan", level=2)
    business_wide = [rec for rec in recommendations if rec.get("finding_type") in _BUSINESS_WIDE_FINDING_TYPES]
    stock_and_products = [rec for rec in recommendations if rec.get("finding_type") not in _BUSINESS_WIDE_FINDING_TYPES]
    if not business_wide and not stock_and_products:
        document.add_paragraph("Nothing to flag for this period.")
    else:
        if business_wide:
            document.add_heading("Business performance", level=3)
            _add_recommendations(document, business_wide, findings)
        if stock_and_products:
            document.add_heading("Stock & products", level=3)
            _add_recommendations(document, stock_and_products, findings)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
